import os
import docx
import json
from docx.shared import Inches


def write_section(lst, text, doc):
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(text)
    run.bold = True
    run.add_break()

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Раздел'
    hdr_cells[1].text = 'Описание'

    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.bold = True

    for row in lst:
        row_cells = table.add_row().cells
        for i in range(len(row)):
            row_cells[i].text = row[i]

    doc.add_page_break()



def make_word_doc(job_id, company, house):

    doc = docx.Document()

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(company)
    run.bold = True
    run.add_break()

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(house)
    run.bold = True
    run.add_break()

    with open(f'static/data/{job_id}/data.json', 'r') as f:
        data = json.load(f)

    photo = data['photo'].copy()
    del data['id']
    del data['photo']
    del data['path']

    for key in data:
        lst = []
        if type(data[key]) ==  str:
            lst.append([key, data[key].replace("'", '').
                        replace('{', '\n').replace('}, ', '\n').replace('}', '').replace('Другое: ', '')])
        else:
            for in_key in data[key]:
                lst.append([in_key, str(data[key][in_key]).replace("'", '').
                            replace('{', '\n').replace('}, ', '\n').replace('}', '').replace('Другое: ', '')])


        write_section(lst, key, doc)
        if key in photo:
            for ph_key in photo[key]:
                for f_name in photo[key][ph_key]:
                    doc.add_picture(os.path.join(f'static/data/{job_id}/photo', f_name), width=Inches(3.5))
                    p = doc.add_paragraph()
                    p.alignment = 1
                    p.add_run(ph_key)
            doc.add_page_break()

    doc.add_page_break()
    doc.save(f'static/data/{job_id}/отчет.docx')
